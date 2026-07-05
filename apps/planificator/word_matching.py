import base64
import binascii
import html
import io
import re
from datetime import date, datetime
from typing import Any

from docx import Document
from rapidfuzz import fuzz

from .constants import ROMANIAN_MONTH_NAMES
from .file_handlers import read_tabular_rows
from .validators import (
    MAX_COURSE_ROWS,
    MAX_TABULAR_COLUMNS,
    MAX_WORD_UPLOAD_BYTES,
    ClientInputError,
)


ENGLISH_MONTH_NAMES = {
    1: "January",
    2: "February",
    3: "March",
    4: "April",
    5: "May",
    6: "June",
    7: "July",
    8: "August",
    9: "September",
    10: "October",
    11: "November",
    12: "December",
}
MIN_MATCH_SCORE = 88
MIN_TOKEN_COVERAGE = 70
MIN_MATCH_GAP = 8
MAX_TEXT_LENGTH = 1_000
MAX_DATE_LENGTH = 50


def normalize_title(value: str) -> str:
    normalized = html.unescape(str(value or "")).replace("\xa0", " ").strip().lower()
    normalized = normalized.translate(
        str.maketrans(
            {
                "ă": "a",
                "â": "a",
                "î": "i",
                "ș": "s",
                "ş": "s",
                "ț": "t",
                "ţ": "t",
            }
        )
    )
    normalized = re.sub(r"[^\w\s]", " ", normalized)
    return re.sub(r"\s+", " ", normalized).strip()


def _cell_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%d.%m.%Y")
    if isinstance(value, date):
        return value.strftime("%d.%m.%Y")
    return str(value).strip()


def read_schedule_rows(file_data: bytes, file_extension: str) -> list[dict[str, Any]]:
    raw_rows = read_tabular_rows(file_data, file_extension.lower())
    if not raw_rows:
        raise ClientInputError("Fișierul de program nu conține un antet.")

    header = [_cell_text(value) for value in raw_rows[0]]
    if len(header) > MAX_TABULAR_COLUMNS:
        raise ClientInputError(f"Fișierul poate avea cel mult {MAX_TABULAR_COLUMNS} coloane.")
    normalized = {name.lower(): index for index, name in enumerate(header) if name}
    if "title" not in normalized:
        raise ClientInputError('Fișierul de program trebuie să conțină coloana "Title".')

    month_indexes: list[int] = []
    for month in range(1, 13):
        aliases = (
            ROMANIAN_MONTH_NAMES[month].lower(),
            ENGLISH_MONTH_NAMES[month].lower(),
        )
        month_index = next((normalized[alias] for alias in aliases if alias in normalized), None)
        if month_index is not None:
            month_indexes.append(month_index)
    if not month_indexes:
        raise ClientInputError(
            "Fișierul de program trebuie să conțină coloane lunare în română sau engleză."
        )

    title_index = normalized["title"]
    schedule_rows: list[dict[str, Any]] = []
    for row_index, raw_row in enumerate(raw_rows[1:]):
        if len(raw_row) > MAX_TABULAR_COLUMNS:
            raise ClientInputError(f"Rândul {row_index + 2} depășește limita de coloane.")
        title = _cell_text(raw_row[title_index]) if title_index < len(raw_row) else ""
        if not title:
            continue
        if len(title) > MAX_TEXT_LENGTH:
            raise ClientInputError(f"Rândul {row_index + 2} conține un titlu prea lung.")
        if len(schedule_rows) >= MAX_COURSE_ROWS:
            raise ClientInputError(
                f"Fișierul poate conține cel mult {MAX_COURSE_ROWS} cursuri.",
                status=413,
            )

        dates: list[str] = []
        for month_index in month_indexes:
            value = _cell_text(raw_row[month_index]) if month_index < len(raw_row) else ""
            if value and value.lower() != "nan":
                if len(value) > MAX_DATE_LENGTH:
                    raise ClientInputError(f"Rândul {row_index + 2} conține o perioadă prea lungă.")
                dates.append(value)
            if len(dates) == 3:
                break
        dates.extend([""] * (3 - len(dates)))
        schedule_rows.append(
            {
                "row_index": row_index,
                "title": title,
                "normalized_title": normalize_title(title),
                "dates": dates,
            }
        )

    if not schedule_rows:
        raise ClientInputError("Fișierul de program nu conține cursuri valide.")
    return schedule_rows


def build_word_course_rows(document: Document) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for table in document.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 6:
                continue
            first_cell = cells[0]
            if first_cell._tc in {cell._tc for cell in cells[1:]}:
                continue
            course_title = first_cell.text.strip()
            if not course_title or sum(1 for cell in cells if cell.text.strip()) <= 1:
                continue
            rows.append(
                {
                    "row": row,
                    "title": course_title,
                    "normalized_title": normalize_title(course_title),
                }
            )
    return rows


def _token_coverage(source: str, target: str) -> float:
    source_tokens = set(source.split())
    target_tokens = set(target.split())
    if not source_tokens:
        return 0
    return len(source_tokens & target_tokens) / len(source_tokens) * 100


def _combined_title_score(word_normalized: str, schedule_normalized: str) -> float:
    return (
        0.35 * fuzz.WRatio(word_normalized, schedule_normalized)
        + 0.25 * fuzz.token_sort_ratio(word_normalized, schedule_normalized)
        + 0.20 * _token_coverage(word_normalized, schedule_normalized)
        + 0.20 * _token_coverage(schedule_normalized, word_normalized)
    )


def _standard_code_tokens(value: str) -> set[str]:
    return set(re.findall(r"\b\d{4,5}\b", value))


def score_schedule_matches(
    word_title: str,
    schedule_rows: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    word_normalized = normalize_title(word_title)
    if not word_normalized or not schedule_rows:
        return []
    exact_matches = [
        entry for entry in schedule_rows if entry["normalized_title"] == word_normalized
    ]
    if len(exact_matches) == 1:
        return [
            {
                "entry": exact_matches[0],
                "score": 100,
                "word_coverage": 100,
                "schedule_coverage": 100,
                "exact": True,
            }
        ]

    scored_matches = []
    for entry in schedule_rows:
        schedule_normalized = entry["normalized_title"]
        scored_matches.append(
            {
                "entry": entry,
                "score": _combined_title_score(word_normalized, schedule_normalized),
                "word_coverage": _token_coverage(word_normalized, schedule_normalized),
                "schedule_coverage": _token_coverage(schedule_normalized, word_normalized),
                "exact": False,
            }
        )
    scored_matches.sort(key=lambda item: item["score"], reverse=True)
    return scored_matches


def confident_match_from_scores(
    word_title: str,
    scored_matches: list[dict[str, Any]],
    match_settings: dict[str, float] | None = None,
) -> dict[str, Any] | None:
    if not scored_matches:
        return None
    if scored_matches[0].get("exact"):
        return scored_matches[0]["entry"]

    match_settings = match_settings or {
        "min_match_score": MIN_MATCH_SCORE,
        "min_token_coverage": MIN_TOKEN_COVERAGE,
        "min_match_gap": MIN_MATCH_GAP,
    }
    word_normalized = normalize_title(word_title)
    best = scored_matches[0]
    second = scored_matches[1] if len(scored_matches) > 1 else None
    below_threshold = (
        best["score"] < float(match_settings["min_match_score"])
        or best["word_coverage"] < float(match_settings["min_token_coverage"])
        or best["schedule_coverage"] < float(match_settings["min_token_coverage"])
    )
    if below_threshold:
        word_codes = _standard_code_tokens(word_normalized)
        best_codes = _standard_code_tokens(best["entry"]["normalized_title"])
        second_codes = (
            _standard_code_tokens(second["entry"]["normalized_title"]) if second else set()
        )
        code_match = (
            word_codes
            and word_codes.issubset(best_codes)
            and not word_codes.issubset(second_codes)
            and best["score"] >= 70
            and best["word_coverage"] >= 60
        )
        if not code_match:
            return None

    if second and best["score"] - second["score"] < float(match_settings["min_match_gap"]):
        word_codes = _standard_code_tokens(word_normalized)
        best_codes = _standard_code_tokens(best["entry"]["normalized_title"])
        second_codes = _standard_code_tokens(second["entry"]["normalized_title"])
        code_match = (
            word_codes
            and word_codes.issubset(best_codes)
            and not word_codes.issubset(second_codes)
            and best["word_coverage"] >= 60
        )
        if not code_match:
            return None
    return best["entry"]


def _candidate_payload(scored_match: dict[str, Any]) -> dict[str, Any]:
    entry = scored_match["entry"]
    return {
        "row_index": entry["row_index"],
        "title": entry["title"],
        "dates": entry["dates"],
        "score": round(scored_match["score"], 1),
        "word_coverage": round(scored_match["word_coverage"], 1),
        "schedule_coverage": round(scored_match["schedule_coverage"], 1),
        "exact": bool(scored_match.get("exact")),
    }


def build_word_match_preview(
    word_file_bytes: bytes,
    schedule_rows: list[dict[str, Any]],
    match_settings: dict[str, Any],
) -> dict[str, Any]:
    document = Document(io.BytesIO(word_file_bytes))
    word_rows = build_word_course_rows(document)
    if not word_rows:
        raise ClientInputError(
            "Documentul Word nu conține rânduri de curs compatibile cu structura așteptată."
        )

    clean_settings = {
        "min_match_score": float(match_settings.get("min_match_score", MIN_MATCH_SCORE)),
        "min_token_coverage": float(
            match_settings.get("min_token_coverage", MIN_TOKEN_COVERAGE)
        ),
        "min_match_gap": float(match_settings.get("min_match_gap", MIN_MATCH_GAP)),
    }
    rows = []
    matched_count = 0
    for word_index, word_row in enumerate(word_rows):
        scored_matches = score_schedule_matches(word_row["title"], schedule_rows)
        confident_match = confident_match_from_scores(
            word_row["title"], scored_matches, clean_settings
        )
        selected_row_index = confident_match["row_index"] if confident_match else None
        matched_count += int(confident_match is not None)
        rows.append(
            {
                "word_row_index": word_index,
                "word_title": word_row["title"],
                "selected_row_index": selected_row_index,
                "status": "matched" if confident_match else "needs_review",
                "candidates": [_candidate_payload(match) for match in scored_matches[:5]],
            }
        )
    return {
        "rows": rows,
        "schedule_options": [
            {
                "row_index": row["row_index"],
                "title": row["title"],
                "dates": row["dates"],
            }
            for row in schedule_rows
        ],
        "matched_count": matched_count,
        "skipped_count": len(rows) - matched_count,
    }


def decode_word_base64(value: Any) -> bytes:
    if not isinstance(value, str) or not value:
        raise ClientInputError("Datele documentului Word lipsesc. Refaceți previzualizarea.")
    if len(value) > (MAX_WORD_UPLOAD_BYTES * 4 // 3) + 8:
        raise ClientInputError("Datele documentului Word depășesc limita permisă.", status=413)
    try:
        decoded = base64.b64decode(value, validate=True)
    except (binascii.Error, ValueError) as exc:
        raise ClientInputError("Datele documentului Word sunt invalide.") from exc
    if not decoded or len(decoded) > MAX_WORD_UPLOAD_BYTES:
        raise ClientInputError("Datele documentului Word au o dimensiune invalidă.", status=413)
    return decoded


def validate_schedule_options(value: Any) -> list[dict[str, Any]]:
    if not isinstance(value, list):
        raise ClientInputError("Opțiunile programului trebuie trimise ca listă.")
    if not value or len(value) > MAX_COURSE_ROWS:
        raise ClientInputError("Numărul opțiunilor programului este invalid.")

    rows = []
    seen_indexes: set[int] = set()
    for item in value:
        if not isinstance(item, dict) or set(item) != {"row_index", "title", "dates"}:
            raise ClientInputError("O opțiune din program are o structură invalidă.")
        row_index = item["row_index"]
        if isinstance(row_index, bool) or not isinstance(row_index, int) or row_index < 0:
            raise ClientInputError("Un index din program este invalid.")
        if row_index in seen_indexes:
            raise ClientInputError("Programul conține indexuri duplicate.")
        seen_indexes.add(row_index)

        title = item["title"]
        dates = item["dates"]
        if not isinstance(title, str) or not title.strip() or len(title.strip()) > MAX_TEXT_LENGTH:
            raise ClientInputError("Un titlu din program este invalid.")
        if not isinstance(dates, list) or len(dates) != 3:
            raise ClientInputError("Fiecare curs trebuie să conțină exact trei perioade.")
        clean_dates = []
        for date_value in dates:
            if not isinstance(date_value, str) or len(date_value.strip()) > MAX_DATE_LENGTH:
                raise ClientInputError("O perioadă din program este invalidă.")
            clean_dates.append(date_value.strip())
        rows.append(
            {
                "row_index": row_index,
                "title": title.strip(),
                "normalized_title": normalize_title(title),
                "dates": clean_dates,
            }
        )
    return rows


def validate_matches(value: Any) -> dict[int, int]:
    if not isinstance(value, list) or len(value) > MAX_COURSE_ROWS:
        raise ClientInputError("Selecțiile trebuie trimise ca listă.")
    matches: dict[int, int] = {}
    for item in value:
        if not isinstance(item, dict) or set(item) != {
            "word_row_index",
            "schedule_row_index",
        }:
            raise ClientInputError("O selecție are o structură invalidă.")
        word_index = item["word_row_index"]
        schedule_index = item["schedule_row_index"]
        if (
            isinstance(word_index, bool)
            or not isinstance(word_index, int)
            or word_index < 0
            or isinstance(schedule_index, bool)
            or not isinstance(schedule_index, int)
            or schedule_index < 0
        ):
            raise ClientInputError("O selecție conține un index invalid.")
        if word_index in matches:
            raise ClientInputError("Selecțiile conțin rânduri Word duplicate.")
        matches[word_index] = schedule_index
    return matches


def apply_word_matches(
    word_file_bytes: bytes,
    schedule_rows: list[dict[str, Any]],
    matches: dict[int, int],
) -> tuple[bytes, int, int]:
    document = Document(io.BytesIO(word_file_bytes))
    word_rows = build_word_course_rows(document)
    schedule_by_index = {row["row_index"]: row for row in schedule_rows}
    if any(index >= len(word_rows) for index in matches):
        raise ClientInputError("O selecție indică un rând Word inexistent.")
    if any(index not in schedule_by_index for index in matches.values()):
        raise ClientInputError("O selecție indică un rând de program inexistent.")

    matched_count = 0
    for word_index, word_row in enumerate(word_rows):
        selected_schedule = schedule_by_index.get(matches.get(word_index))
        if not selected_schedule:
            continue
        for target_index, date_value in zip((3, 4, 5), selected_schedule["dates"]):
            word_row["row"].cells[target_index].text = date_value
        matched_count += 1

    output = io.BytesIO()
    document.save(output)
    return output.getvalue(), matched_count, len(word_rows) - matched_count
